import streamlit as st
import pandas as pd
import re
import numpy as np
import matplotlib.pyplot as plt
import geopandas as gpd

# Custom CSS with blue and white theme and zoom functionality
st.markdown("""
<style>
    /* Allow zoom functionality */
    .stApp {
        zoom: 1 !important;
        transform: scale(1) !important;
        transform-origin: 0 0 !important;
    }
    
    /* Increase sidebar width */
    section[data-testid="stSidebar"] {
        width: 320px !important;
        background-color: #f8f9fd !important;
    }
    
    /* Main app styling with blue theme */
    .main .block-container {
        padding-left: 1rem !important;
        padding-right: 1rem !important;
        margin-left: 0 !important;
        max-width: none !important;
        background-color: white !important;
        border-radius: 10px !important;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1) !important;
        margin-top: 1rem !important;
        margin-bottom: 1rem !important;
    }
    
    /* Title styling */
    h1 {
        color: #2c3e50 !important;
        text-align: center !important;
        font-weight: 700 !important;
        margin-bottom: 2rem !important;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1) !important;
    }
    
    /* Subheader styling */
    h2, h3 {
        color: #34495e !important;
        border-bottom: 2px solid #3498db !important;
        padding-bottom: 0.5rem !important;
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        padding-left: 1rem !important;
        margin-left: 0 !important;
    }
    
    /* Button styling */
    .stButton > button {
        background: linear-gradient(45deg, #3498db, #2980b9) !important;
        color: white !important;
        border: none !important;
        border-radius: 25px !important;
        padding: 0.5rem 2rem !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 15px rgba(52, 152, 219, 0.3) !important;
        transition: all 0.3s ease !important;
    }
    
    .stButton > button:hover {
        background: linear-gradient(45deg, #2980b9, #1f5f8b) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(52, 152, 219, 0.4) !important;
    }
    
    /* Download button styling */
    .stDownloadButton > button {
        background: linear-gradient(45deg, #27ae60, #229954) !important;
        color: white !important;
        border: none !important;
        border-radius: 25px !important;
        padding: 0.5rem 2rem !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 15px rgba(39, 174, 96, 0.3) !important;
        transition: all 0.3s ease !important;
    }
    
    .stDownloadButton > button:hover {
        background: linear-gradient(45deg, #229954, #1e7e34) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(39, 174, 96, 0.4) !important;
    }
    
    /* Metric styling */
    [data-testid="metric-container"] {
        background: linear-gradient(135deg, #74b9ff, #0984e3) !important;
        border: 1px solid #ddd !important;
        padding: 1rem !important;
        border-radius: 10px !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
    }
    
    [data-testid="metric-container"] > div {
        color: white !important;
    }
    
    /* Dataframe styling */
    .stDataFrame {
        border-radius: 10px !important;
        overflow: hidden !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
    }
    
    /* Info box styling */
    .stInfo {
        background: linear-gradient(135deg, #74b9ff, #0984e3) !important;
        color: white !important;
        border-radius: 10px !important;
    }
    
    /* Warning box styling */
    .stWarning {
        background: linear-gradient(135deg, #fdcb6e, #e17055) !important;
        color: white !important;
        border-radius: 10px !important;
    }
    
    /* Success box styling */
    .stSuccess {
        background: linear-gradient(135deg, #00b894, #00a085) !important;
        color: white !important;
        border-radius: 10px !important;
    }
    
    /* Logo container styling for consistent alignment */
    .logo-container {
        display: flex !important;
        justify-content: space-between !important;
        align-items: center !important;
        padding: 1rem 0 !important;
        margin-bottom: 1rem !important;
        border-bottom: 2px solid #f0f0f0 !important;
    }
    
    .logo-item {
        display: flex !important;
        justify-content: center !important;
        align-items: center !important;
        width: 300px !important;
        height: 200px !important;
        border: 2px solid #3498db !important;
        border-radius: 15px !important;
        background: linear-gradient(135deg, #f8f9fd, #e3f2fd) !important;
        overflow: hidden !important;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1) !important;
    }
    
    .logo-item img {
        max-width: 280px !important;
        max-height: 180px !important;
        width: auto !important;
        height: auto !important;
        object-fit: contain !important;
    }
    
    .logo-placeholder {
        width: 280px !important;
        height: 180px !important;
        border: 2px dashed #3498db !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        background: linear-gradient(135deg, #f8f9fd, #e3f2fd) !important;
        border-radius: 15px !important;
        font-size: 14px !important;
        color: #2c3e50 !important;
        text-align: center !important;
        font-weight: 600 !important;
    }
    
    /* Report export button styling */
    .report-button {
        background: linear-gradient(45deg, #e74c3c, #c0392b) !important;
        color: white !important;
        border: none !important;
        border-radius: 25px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 700 !important;
        font-size: 16px !important;
        box-shadow: 0 4px 15px rgba(231, 76, 60, 0.3) !important;
        transition: all 0.3s ease !important;
        cursor: pointer !important;
        width: 100% !important;
        margin: 1rem 0 !important;
    }
    
    .report-button:hover {
        background: linear-gradient(45deg, #c0392b, #a93226) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(231, 76, 60, 0.4) !important;
    }
</style>
""", unsafe_allow_html=True)

# Logo Section with consistent alignment
st.markdown("""
<div class="logo-container">
    <div class="logo-item" id="logo1"></div>
    <div class="logo-item" id="logo2"></div>
    <div class="logo-item" id="logo3"></div>
</div>
""", unsafe_allow_html=True)

# Add logos with consistent sizing using columns
col1, col2, col3 = st.columns(3)
with col1:
    try:
        st.markdown('<div style="text-align: center; margin-bottom: 1rem;">', unsafe_allow_html=True)
        st.image("NMCP.png", width=280, caption="National Malaria Control Program")
        st.markdown('</div>', unsafe_allow_html=True)
    except:
        st.markdown("""
        <div class="logo-placeholder">
            NMCP.png<br>
            Not Found<br>
            Please upload logo
        </div>
        """, unsafe_allow_html=True)

with col2:
    try:
        st.markdown('<div style="text-align: center; margin-bottom: 1rem;">', unsafe_allow_html=True)
        st.image("icf_sl (2).jpg", width=280, caption="ICF Sierra Leone")
        st.markdown('</div>', unsafe_allow_html=True)
    except:
        st.markdown("""
        <div class="logo-placeholder">
            icf_sl (2).jpg<br>
            Not Found<br>
            Please upload logo
        </div>
        """, unsafe_allow_html=True)

with col3:
    # Placeholder for third logo
    st.markdown("""
    <div class="logo-placeholder">
        Partner Logo<br>
        (Right Position)<br>
        Upload .png/.jpg
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")  # Add a horizontal line separator

# Streamlit App
st.title("📊 Text Data Extraction & Visualization")

# Upload file
uploaded_file = "latest_sbd_06_10_2025.xlsx"
if uploaded_file:
    # Read the uploaded Excel file
    df_original = pd.read_excel(uploaded_file)
    
    # Load shapefile
    try:
        gdf = gpd.read_file("Chiefdom 2021.shp")
        st.success("✅ Shapefile loaded successfully!")
    except Exception as e:
        st.error(f"❌ Could not load shapefile: {e}")
        gdf = None
    
    # Create empty lists to store extracted data
    districts, chiefdoms, phu_names, community_names, school_names = [], [], [], [], []
    
    # Process each row in the "Scan QR code" column
    for qr_text in df_original["Scan QR code"]:
        if pd.isna(qr_text):
            districts.append(None)
            chiefdoms.append(None)
            phu_names.append(None)
            community_names.append(None)
            school_names.append(None)
            continue
            
        # Extract values using regex patterns
        district_match = re.search(r"District:\s*([^\n]+)", str(qr_text))
        districts.append(district_match.group(1).strip() if district_match else None)
        
        chiefdom_match = re.search(r"Chiefdom:\s*([^\n]+)", str(qr_text))
        chiefdoms.append(chiefdom_match.group(1).strip() if chiefdom_match else None)
        
        phu_match = re.search(r"PHU name:\s*([^\n]+)", str(qr_text))
        phu_names.append(phu_match.group(1).strip() if phu_match else None)
        
        community_match = re.search(r"Community name:\s*([^\n]+)", str(qr_text))
        community_names.append(community_match.group(1).strip() if community_match else None)
        
        school_match = re.search(r"Name of school:\s*([^\n]+)", str(qr_text))
        school_names.append(school_match.group(1).strip() if school_match else None)
    
    # Create a new DataFrame with extracted values
    extracted_df = pd.DataFrame({
        "District": districts,
        "Chiefdom": chiefdoms,
        "PHU Name": phu_names,
        "Community Name": community_names,
        "School Name": school_names
    })
    
    # Add all other columns from the original DataFrame
    for column in df_original.columns:
        if column != "Scan QR code":  # Skip the QR code column since we've already processed it
            extracted_df[column] = df_original[column]
    
    # Create sidebar filters early so they're available for all sections
    st.sidebar.header("Filter Options")
    
    # Create radio buttons to select which level to group by
    grouping_selection = st.sidebar.radio(
        "Select the level for grouping:",
        ["District", "Chiefdom", "PHU Name", "Community Name", "School Name"],
        index=0  # Default to 'District'
    )
    
    # Dictionary to define the hierarchy for each grouping level
    hierarchy = {
        "District": ["District"],
        "Chiefdom": ["District", "Chiefdom"],
        "PHU Name": ["District", "Chiefdom", "PHU Name"],
        "Community Name": ["District", "Chiefdom", "PHU Name", "Community Name"],
        "School Name": ["District", "Chiefdom", "PHU Name", "Community Name", "School Name"]
    }
    
    # Initialize filtered dataframe with the full dataset
    filtered_df = extracted_df.copy()
    
    # Dictionary to store selected values for each level
    selected_values = {}
    
    # Apply filters based on the hierarchy for the selected grouping level
    for level in hierarchy[grouping_selection]:
        # Filter out None/NaN values and get sorted unique values
        level_values = sorted(filtered_df[level].dropna().unique())
        
        if level_values:
            # Create selectbox for this level
            selected_value = st.sidebar.selectbox(f"Select {level}", level_values)
            selected_values[level] = selected_value
            
            # Apply filter to the dataframe
            filtered_df = filtered_df[filtered_df[level] == selected_value]
    
    # Display Dual Maps at the top
    st.subheader("🗺️ Geographic Distribution Maps")
    
    if gdf is not None:
        # Define specific districts for left and right maps
        left_district = "BO"
        right_district = "BOMBALI"
        
        # BO DISTRICT MAP - Full width
        st.write(f"**{left_district} District - All Chiefdoms**")
        
        # Filter shapefile for BO district
        bo_gdf = gdf[gdf['FIRST_DNAM'] == left_district].copy()
        
        if len(bo_gdf) > 0:
            # Filter data for BO district to get GPS coordinates
            bo_data = extracted_df[extracted_df["District"] == left_district].copy()
            
            # Create the BO district plot
            fig_bo, ax_bo = plt.subplots(figsize=(14, 8))
            
            # Plot chiefdom boundaries in white with black edges
            bo_gdf.plot(ax=ax_bo, color='white', edgecolor='black', alpha=0.8, linewidth=1)
            
            # Add chiefdom labels
            for idx, row in bo_gdf.iterrows():
                if 'FIRST_CHIE' in row and pd.notna(row['FIRST_CHIE']):
                    centroid = row.geometry.centroid
                    ax_bo.annotate(
                        row['FIRST_CHIE'], 
                        (centroid.x, centroid.y),
                        xytext=(5, 5), 
                        textcoords='offset points',
                        fontsize=10,
                        ha='left',
                        bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.8)
                    )
            
            # Plot GPS coordinates if available
            if len(bo_data) > 0 and "GPS Location" in bo_data.columns:
                gps_data = bo_data["GPS Location"].dropna()
                coords_extracted = []
                
                for gps_val in gps_data:
                    coords = re.findall(r'-?\d+\.?\d*', str(gps_val))
                    if len(coords) >= 2:
                        try:
                            lat, lon = float(coords[0]), float(coords[1])
                            coords_extracted.append([lat, lon])
                        except:
                            continue
                
                if coords_extracted:
                    lats, lons = zip(*coords_extracted)
                    ax_bo.scatter(
                        lons, lats,
                        c='blue',
                        s=60,
                        alpha=0.8,
                        edgecolors='darkblue',
                        linewidth=1,
                        zorder=5,
                        label=f'Schools ({len(coords_extracted)})'
                    )
                    ax_bo.legend()
            
            # Customize plot
            ax_bo.set_title(f'{left_district} District - Chiefdoms: {len(bo_gdf)}', fontsize=16, fontweight='bold')
            ax_bo.set_xlabel('Longitude', fontsize=12)
            ax_bo.set_ylabel('Latitude', fontsize=12)
            
            # Remove axis ticks for cleaner look
            ax_bo.set_xticks([])
            ax_bo.set_yticks([])
            
            plt.tight_layout()
            st.pyplot(fig_bo)
            
            # Display chiefdoms list
            if 'FIRST_CHIE' in bo_gdf.columns:
                chiefdoms = bo_gdf['FIRST_CHIE'].dropna().tolist()
                st.write(f"**Chiefdoms in {left_district} District ({len(chiefdoms)}):**")
                chiefdom_cols = st.columns(3)
                for i, chiefdom in enumerate(chiefdoms):
                    with chiefdom_cols[i % 3]:
                        st.write(f"• {chiefdom}")
        else:
            st.warning(f"No chiefdoms found for {left_district} district in shapefile")
        
        st.divider()
        
        # BOMBALI DISTRICT MAP - Full width
        st.write(f"**{right_district} District - All Chiefdoms**")
        
        # Filter shapefile for BOMBALI district
        bombali_gdf = gdf[gdf['FIRST_DNAM'] == right_district].copy()
        
        if len(bombali_gdf) > 0:
            # Filter data for BOMBALI district to get GPS coordinates
            bombali_data = extracted_df[extracted_df["District"] == right_district].copy()
            
            # Create the BOMBALI district plot
            fig_bombali, ax_bombali = plt.subplots(figsize=(14, 8))
            
            # Plot chiefdom boundaries in white with black edges
            bombali_gdf.plot(ax=ax_bombali, color='white', edgecolor='black', alpha=0.8, linewidth=1)
            
            # Add chiefdom labels
            for idx, row in bombali_gdf.iterrows():
                if 'FIRST_CHIE' in row and pd.notna(row['FIRST_CHIE']):
                    centroid = row.geometry.centroid
                    ax_bombali.annotate(
                        row['FIRST_CHIE'], 
                        (centroid.x, centroid.y),
                        xytext=(5, 5), 
                        textcoords='offset points',
                        fontsize=10,
                        ha='left',
                        bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.8)
                    )
            
            # Plot GPS coordinates if available
            if len(bombali_data) > 0 and "GPS Location" in bombali_data.columns:
                gps_data = bombali_data["GPS Location"].dropna()
                coords_extracted = []
                
                for gps_val in gps_data:
                    coords = re.findall(r'-?\d+\.?\d*', str(gps_val))
                    if len(coords) >= 2:
                        try:
                            lat, lon = float(coords[0]), float(coords[1])
                            coords_extracted.append([lat, lon])
                        except:
                            continue
                
                if coords_extracted:
                    lats, lons = zip(*coords_extracted)
                    ax_bombali.scatter(
                        lons, lats,
                        c='blue',
                        s=60,
                        alpha=0.8,
                        edgecolors='darkblue',
                        linewidth=1,
                        zorder=5,
                        label=f'Schools ({len(coords_extracted)})'
                    )
                    ax_bombali.legend()
            
            # Customize plot
            ax_bombali.set_title(f'{right_district} District - Chiefdoms: {len(bombali_gdf)}', fontsize=16, fontweight='bold')
            ax_bombali.set_xlabel('Longitude', fontsize=12)
            ax_bombali.set_ylabel('Latitude', fontsize=12)
            
            # Remove axis ticks for cleaner look
            ax_bombali.set_xticks([])
            ax_bombali.set_yticks([])
            
            plt.tight_layout()
            st.pyplot(fig_bombali)
            
            # Display chiefdoms list
            if 'FIRST_CHIE' in bombali_gdf.columns:
                chiefdoms = bombali_gdf['FIRST_CHIE'].dropna().tolist()
                st.write(f"**Chiefdoms in {right_district} District ({len(chiefdoms)}):**")
                chiefdom_cols = st.columns(3)
                for i, chiefdom in enumerate(chiefdoms):
                    with chiefdom_cols[i % 3]:
                        st.write(f"• {chiefdom}")
        else:
            st.warning(f"No chiefdoms found for {right_district} district in shapefile")
    else:
        st.error("Shapefile not loaded. Cannot display map.")
    
    # Display Original Data Sample
    st.subheader("📄 Original Data Sample")
    st.dataframe(df_original.head())
    
    # Display Extracted Data
    st.subheader("📋 Extracted Data")
    st.dataframe(extracted_df)
    
    # Add download button for CSV
    csv = extracted_df.to_csv(index=False)
    st.download_button(
        label="📥 Download Extracted Data as CSV",
        data=csv,
        file_name="extracted_school_data.csv",
        mime="text/csv"
    )
    
    # Enrollment and ITN Distribution Analysis
    st.subheader("📊 Enrollment and ITN Distribution Analysis")
    
    # Calculate total enrollment and ITN distribution by district
    district_analysis = []
    
    for district in extracted_df['District'].dropna().unique():
        district_data = extracted_df[extracted_df['District'] == district]
        
        total_enrollment = 0
        total_itn = 0
        
        # Sum enrollments and ITN by class
        for class_num in range(1, 6):
            boys_col = f"Number of boys in class {class_num}"
            girls_col = f"Number of girls in class {class_num}"
            itn_col = f"Number of ITN distributed to class {class_num}"
            
            if boys_col in district_data.columns:
                total_enrollment += district_data[boys_col].fillna(0).sum()
            if girls_col in district_data.columns:
                total_enrollment += district_data[girls_col].fillna(0).sum()
            if itn_col in district_data.columns:
                total_itn += district_data[itn_col].fillna(0).sum()
        
        # Calculate coverage percentage
        coverage = (total_itn / total_enrollment * 100) if total_enrollment > 0 else 0
        
        district_analysis.append({
            'District': district,
            'Total_Enrollment': total_enrollment,
            'Total_ITN': total_itn,
            'Coverage': coverage
        })
    
    district_df = pd.DataFrame(district_analysis)
    
    # Create bar charts for district analysis
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
    
    # Total Enrollment by District
    ax1.bar(district_df['District'], district_df['Total_Enrollment'], color='skyblue', edgecolor='navy')
    ax1.set_title('Total Enrollment by District', fontsize=14, fontweight='bold')
    ax1.set_ylabel('Number of Students')
    ax1.tick_params(axis='x', rotation=45)
    
    # Add value labels on bars
    for i, v in enumerate(district_df['Total_Enrollment']):
        ax1.text(i, v + max(district_df['Total_Enrollment']) * 0.01, str(int(v)), ha='center', fontweight='bold')
    
    # Total ITN Distributed by District
    ax2.bar(district_df['District'], district_df['Total_ITN'], color='lightcoral', edgecolor='darkred')
    ax2.set_title('Total ITN Distributed by District', fontsize=14, fontweight='bold')
    ax2.set_ylabel('Number of ITNs')
    ax2.tick_params(axis='x', rotation=45)
    
    # Add value labels on bars
    for i, v in enumerate(district_df['Total_ITN']):
        ax2.text(i, v + max(district_df['Total_ITN']) * 0.01, str(int(v)), ha='center', fontweight='bold')
    
    plt.tight_layout()
    st.pyplot(fig)
    
    # District-level pie charts
    st.subheader("📊 District-Level Distribution (Pie Charts)")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Enrollment pie chart
        if district_df['Total_Enrollment'].sum() > 0:
            fig_pie1, ax_pie1 = plt.subplots(figsize=(8, 8))
            colors_enrollment = ['#87CEEB', '#4682B4', '#1E90FF', '#0000CD', '#000080']
            # Filter out zero values for pie chart
            enrollment_data = district_df[district_df['Total_Enrollment'] > 0]
            if len(enrollment_data) > 0:
                wedges, texts, autotexts = ax_pie1.pie(enrollment_data['Total_Enrollment'], 
                                                      labels=enrollment_data['District'],
                                                      autopct='%1.1f%%',
                                                      colors=colors_enrollment[:len(enrollment_data)],
                                                      startangle=90)
                ax_pie1.set_title('Total Enrollment Distribution\nby District', fontsize=14, fontweight='bold')
                plt.setp(autotexts, size=10, weight="bold")
                st.pyplot(fig_pie1)
            else:
                st.warning("No enrollment data available for pie chart")
        else:
            st.warning("No enrollment data available for pie chart")
    
    with col2:
        # ITN distribution pie chart
        if district_df['Total_ITN'].sum() > 0:
            fig_pie2, ax_pie2 = plt.subplots(figsize=(8, 8))
            colors_itn = ['#90EE90', '#32CD32', '#228B22', '#006400', '#004000']
            # Filter out zero values for pie chart
            itn_data = district_df[district_df['Total_ITN'] > 0]
            if len(itn_data) > 0:
                wedges, texts, autotexts = ax_pie2.pie(itn_data['Total_ITN'], 
                                                      labels=itn_data['District'],
                                                      autopct='%1.1f%%',
                                                      colors=colors_itn[:len(itn_data)],
                                                      startangle=90)
                ax_pie2.set_title('Total ITN Distribution\nby District', fontsize=14, fontweight='bold')
                plt.setp(autotexts, size=10, weight="bold")
                st.pyplot(fig_pie2)
            else:
                st.warning("No ITN distribution data available for pie chart")
        else:
            st.warning("No ITN distribution data available for pie chart")
    
    # Chiefdoms Analysis by District
    st.subheader("📊 Chiefdoms Analysis by District")
    
    # Get all unique districts that have chiefdom data
    districts_with_chiefdoms = extracted_df[extracted_df['Chiefdom'].notna()]['District'].unique()
    
    for district in districts_with_chiefdoms:
        st.write(f"### {district} District - Chiefdoms Analysis")
        
        # Filter data for this district
        district_data = extracted_df[extracted_df['District'] == district]
        district_chiefdoms = district_data['Chiefdom'].dropna().unique()
        
        if len(district_chiefdoms) > 0:
            # Calculate by chiefdom for this district
            district_chiefdom_analysis = []
            
            for chiefdom in district_chiefdoms:
                chiefdom_data = district_data[district_data['Chiefdom'] == chiefdom]
                
                total_enrollment = 0
                total_itn = 0
                
                for class_num in range(1, 6):
                    boys_col = f"Number of boys in class {class_num}"
                    girls_col = f"Number of girls in class {class_num}"
                    itn_col = f"Number of ITN distributed to class {class_num}"
                    
                    if boys_col in chiefdom_data.columns:
                        total_enrollment += chiefdom_data[boys_col].fillna(0).sum()
                    if girls_col in chiefdom_data.columns:
                        total_enrollment += chiefdom_data[girls_col].fillna(0).sum()
                    if itn_col in chiefdom_data.columns:
                        total_itn += chiefdom_data[itn_col].fillna(0).sum()
                
                # Calculate coverage percentage
                coverage = (total_itn / total_enrollment * 100) if total_enrollment > 0 else 0
                
                district_chiefdom_analysis.append({
                    'Chiefdom': chiefdom,
                    'Total_Enrollment': total_enrollment,
                    'Total_ITN': total_itn,
                    'Coverage': coverage
                })
            
            district_chiefdom_df = pd.DataFrame(district_chiefdom_analysis)
            district_chiefdom_df = district_chiefdom_df.sort_values('Total_Enrollment', ascending=False)
            
            if len(district_chiefdom_df) > 0:
                # Create 3-column subplot for this district's chiefdoms
                fig_district, (ax1_dist, ax2_dist, ax3_dist) = plt.subplots(1, 3, figsize=(24, 8))
                
                # Total Enrollment by Chiefdoms in this District (Blue)
                bars1 = ax1_dist.barh(district_chiefdom_df['Chiefdom'], district_chiefdom_df['Total_Enrollment'], 
                                      color='#4682B4', edgecolor='navy')
                ax1_dist.set_title(f'{district} District - Total Enrollment by Chiefdom', fontsize=14, fontweight='bold')
                ax1_dist.set_xlabel('Number of Students')
                
                # Add value labels
                for i, v in enumerate(district_chiefdom_df['Total_Enrollment']):
                    if v > 0:  # Only show label if value is greater than 0
                        ax1_dist.text(v + max(district_chiefdom_df['Total_Enrollment']) * 0.01, i, 
                                      str(int(v)), va='center', fontweight='bold')
                
                # Total ITN by Chiefdoms in this District (Green)
                bars2 = ax2_dist.barh(district_chiefdom_df['Chiefdom'], district_chiefdom_df['Total_ITN'], 
                                      color='#32CD32', edgecolor='darkgreen')
                ax2_dist.set_title(f'{district} District - Total ITN Distributed by Chiefdom', fontsize=14, fontweight='bold')
                ax2_dist.set_xlabel('Number of ITNs')
                
                # Add value labels
                for i, v in enumerate(district_chiefdom_df['Total_ITN']):
                    if v > 0:  # Only show label if value is greater than 0
                        ax2_dist.text(v + max(district_chiefdom_df['Total_ITN']) * 0.01, i, 
                                      str(int(v)), va='center', fontweight='bold')
                
                # Coverage by Chiefdoms in this District (Yellow)
                bars3 = ax3_dist.barh(district_chiefdom_df['Chiefdom'], district_chiefdom_df['Coverage'], 
                                      color='#FFD700', edgecolor='orange')
                ax3_dist.set_title(f'{district} District - ITN Coverage by Chiefdom (%)', fontsize=14, fontweight='bold')
                ax3_dist.set_xlabel('Coverage Percentage (%)')
                
                # Add value labels
                for i, v in enumerate(district_chiefdom_df['Coverage']):
                    if v > 0:  # Only show label if value is greater than 0
                        ax3_dist.text(v + max(district_chiefdom_df['Coverage']) * 0.01, i, 
                                      f'{v:.1f}%', va='center', fontweight='bold')
                
                plt.tight_layout()
                st.pyplot(fig_district)
                
                # Display summary table for this district
                st.write(f"**{district} District Summary:**")
                summary_cols = st.columns(3)
                with summary_cols[0]:
                    st.metric("Total Chiefdoms", len(district_chiefdom_df))
                with summary_cols[1]:
                    st.metric("Total Students", int(district_chiefdom_df['Total_Enrollment'].sum()))
                with summary_cols[2]:
                    st.metric("Total ITNs", int(district_chiefdom_df['Total_ITN'].sum()))
                
                st.divider()
            else:
                st.warning(f"No chiefdom data available for {district} district")
        else:
            st.warning(f"No chiefdoms found for {district} district")
    
    # Summary buttons section
    st.subheader("📊 Summary Reports")
    
    # Create two columns for the summary buttons
    col1, col2 = st.columns(2)
    
    # Button for District Summary
    with col1:
        district_summary_button = st.button("Show District Summary")
    
    # Button for Chiefdom Summary
    with col2:
        chiefdom_summary_button = st.button("Show Chiefdom Summary")
    
    # Display District Summary when button is clicked
    if district_summary_button:
        st.subheader("📈 Summary by District")
        
        # Create aggregation dictionary
        agg_dict = {}
        
        # Add enrollment columns to aggregation
        for class_num in range(1, 6):
            total_col = f"Number of enrollments in class {class_num}"
            boys_col = f"Number of boys in class {class_num}"
            girls_col = f"Number of girls in class {class_num}"
            
            if total_col in extracted_df.columns:
                agg_dict[total_col] = "sum"
            if boys_col in extracted_df.columns:
                agg_dict[boys_col] = "sum"
            if girls_col in extracted_df.columns:
                agg_dict[girls_col] = "sum"
        
        # Group by District and aggregate
        district_summary = extracted_df.groupby("District").agg(agg_dict).reset_index()
        
        # Calculate total enrollment
        district_summary["Total Enrollment"] = 0
        for class_num in range(1, 6):
            total_col = f"Number of enrollments in class {class_num}"
            if total_col in district_summary.columns:
                district_summary["Total Enrollment"] += district_summary[total_col]
        
        # Display summary table
        st.dataframe(district_summary)
        
        # Download button for district summary
        district_csv = district_summary.to_csv(index=False)
        st.download_button(
            label="📥 Download District Summary as CSV",
            data=district_csv,
            file_name="district_summary.csv",
            mime="text/csv"
        )
        
        # Create a bar chart for district summary
        fig, ax = plt.subplots(figsize=(12, 8))
        district_summary.plot(kind="bar", x="District", y="Total Enrollment", ax=ax, color="blue")
        ax.set_title("📊 Total Enrollment by District")
        ax.set_xlabel("")
        ax.set_ylabel("Number of Students")
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        st.pyplot(fig)
    
    # Display Chiefdom Summary when button is clicked
    if chiefdom_summary_button:
        st.subheader("📈 Summary by Chiefdom")
        
        # Create aggregation dictionary
        agg_dict = {}
        
        # Add enrollment columns to aggregation
        for class_num in range(1, 6):
            total_col = f"Number of enrollments in class {class_num}"
            boys_col = f"Number of boys in class {class_num}"
            girls_col = f"Number of girls in class {class_num}"
            
            if total_col in extracted_df.columns:
                agg_dict[total_col] = "sum"
            if boys_col in extracted_df.columns:
                agg_dict[boys_col] = "sum"
            if girls_col in extracted_df.columns:
                agg_dict[girls_col] = "sum"
        
        # Group by District and Chiefdom and aggregate
        chiefdom_summary = extracted_df.groupby(["District", "Chiefdom"]).agg(agg_dict).reset_index()
        
        # Calculate total enrollment
        chiefdom_summary["Total Enrollment"] = 0
        for class_num in range(1, 6):
            total_col = f"Number of enrollments in class {class_num}"
            if total_col in chiefdom_summary.columns:
                chiefdom_summary["Total Enrollment"] += chiefdom_summary[total_col]
        
        # Display summary table
        st.dataframe(chiefdom_summary)
        
        # Download button for chiefdom summary
        chiefdom_csv = chiefdom_summary.to_csv(index=False)
        st.download_button(
            label="📥 Download Chiefdom Summary as CSV",
            data=chiefdom_csv,
            file_name="chiefdom_summary.csv",
            mime="text/csv"
        )
        
        # Create a temporary label for the chart
        chiefdom_summary['Label'] = chiefdom_summary['District'] + '\n' + chiefdom_summary['Chiefdom']
        
        # Create a bar chart for chiefdom summary
        fig, ax = plt.subplots(figsize=(14, 10))
        chiefdom_summary.plot(kind="barh", x="Label", y="Total Enrollment", ax=ax, color="blue")
        ax.set_title("📊 Total Enrollment by District and Chiefdom")
        ax.set_ylabel("")
        ax.set_xlabel("Number of Students")
        plt.tight_layout()
        st.pyplot(fig)
    
    # Visualization and filtering section
    st.subheader("🔍 Detailed Data Filtering and Visualization")
    
    # Check if data is available after filtering
    if not filtered_df.empty:
        st.write(f"### Filtered Data - {len(filtered_df)} records")
        st.dataframe(filtered_df)
        
        # Download button for filtered data
        filtered_csv = filtered_df.to_csv(index=False)
        st.download_button(
            label="📥 Download Filtered Data as CSV",
            data=filtered_csv,
            file_name="filtered_data.csv",
            mime="text/csv"
        )
        
        # Define the hierarchy levels to include in the summary
        group_columns = hierarchy[grouping_selection]
        
        # Create aggregation dictionary for enrollment data
        agg_dict = {}
        for class_num in range(1, 6):
            total_col = f"Number of enrollments in class {class_num}"
            boys_col = f"Number of boys in class {class_num}"
            girls_col = f"Number of girls in class {class_num}"
            
            if total_col in filtered_df.columns:
                agg_dict[total_col] = "sum"
            if boys_col in filtered_df.columns:
                agg_dict[boys_col] = "sum"
            if girls_col in filtered_df.columns:
                agg_dict[girls_col] = "sum"
        
        # Group by the selected hierarchical columns
        grouped_data = filtered_df.groupby(group_columns).agg(agg_dict).reset_index()
        
        # Calculate total enrollment
        grouped_data["Total Enrollment"] = 0
        for class_num in range(1, 6):
            total_col = f"Number of enrollments in class {class_num}"
            if total_col in grouped_data.columns:
                grouped_data["Total Enrollment"] += grouped_data[total_col]
        
        # Summary Table with separate columns for each level
        st.subheader("📊 Detailed Summary Table")
        st.dataframe(grouped_data)
        
        # Create a temporary group column for the chart
        grouped_data['Group'] = grouped_data[group_columns].apply(lambda row: ','.join(row.astype(str)), axis=1)
        
        # Create a bar chart
        fig, ax = plt.subplots(figsize=(12, 8))
        grouped_data.plot(kind="bar", x="Group", y="Total Enrollment", ax=ax, color="blue")
        ax.set_title(f"Total Enrollment by {grouping_selection}")
        
        # Remove x-label as requested
        ax.set_xlabel("")
        
        ax.set_ylabel("Number of Students")
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        st.pyplot(fig)
    else:
        st.warning("No data available for the selected filters.")

    # Final Data Export Section
    st.subheader("📥 Export Complete Dataset")
    st.write("Download the complete extracted dataset in your preferred format:")

    # Create download buttons in columns
    download_col1, download_col2, download_col3 = st.columns(3)

    with download_col1:
        # CSV Download
        csv_data = extracted_df.to_csv(index=False)
        st.download_button(
            label="📄 Download Complete Data as CSV",
            data=csv_data,
            file_name="complete_extracted_data.csv",
            mime="text/csv",
            help="Download all extracted data in CSV format"
        )

    with download_col2:
        # Excel Download
        from io import BytesIO
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            extracted_df.to_excel(writer, sheet_name='Extracted Data', index=False)
        excel_data = excel_buffer.getvalue()
        
        st.download_button(
            label="📊 Download Complete Data as Excel",
            data=excel_data,
            file_name="complete_extracted_data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            help="Download all extracted data in Excel format"
        )

    with download_col3:
        # Word Report Download
        if st.button("📋 Generate Word Report", help="Generate and download comprehensive report in Word format"):
            # Generate Word report content
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from datetime import datetime
            import matplotlib.pyplot as plt
            import io
            
            doc = Document()
            
            # Add logos to header (if available)
            try:
                # Create header section with logos
                header_para = doc.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Try to add logos
                try:
                    logo_run1 = header_para.add_run()
                    logo_run1.add_picture("NMCP.png", width=Inches(1.5))
                    header_para.add_run("    ")  # Space between logos
                except:
                    header_para.add_run("NMCP    ")
                
                try:
                    logo_run2 = header_para.add_run()
                    logo_run2.add_picture("icf_sl (2).jpg", width=Inches(1.5))
                    header_para.add_run("    ")  # Space between logos
                except:
                    header_para.add_run("ICF Sierra Leone    ")
                
                header_para.add_run("Partner Logo")
                
                doc.add_paragraph()  # Add space after logos
            except:
                # If logos fail, add text headers
                header_para = doc.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header_para.add_run("NMCP | ICF Sierra Leone | Partner Organization")
                header_run.font.size = Pt(12)
                header_run.bold = True
            
            # Add title page
            title = doc.add_heading('School-Based Distribution (SBD)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Comprehensive Analysis Report', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date and time
            current_datetime = datetime.now()
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f"Generated on: {current_datetime.strftime('%B %d, %Y at %I:%M %p')}")
            date_run.font.size = Pt(12)
            date_run.bold = True
            
            # Add page break
            doc.add_page_break()
            
            # Add executive summary
            doc.add_heading('Executive Summary', level=1)
            total_records = len(extracted_df)
            total_districts = len(extracted_df['District'].dropna().unique())
            total_chiefdoms = len(extracted_df['Chiefdom'].dropna().unique())
            
            # Calculate overall statistics
            total_enrollment = 0
            total_itn = 0
            for class_num in range(1, 6):
                boys_col = f"Number of boys in class {class_num}"
                girls_col = f"Number of girls in class {class_num}"
                itn_col = f"Number of ITN distributed to class {class_num}"
                
                if boys_col in extracted_df.columns:
                    total_enrollment += extracted_df[boys_col].fillna(0).sum()
                if girls_col in extracted_df.columns:
                    total_enrollment += extracted_df[girls_col].fillna(0).sum()
                if itn_col in extracted_df.columns:
                    total_itn += extracted_df[itn_col].fillna(0).sum()
            
            overall_coverage = (total_itn / total_enrollment * 100) if total_enrollment > 0 else 0
            
            summary_text = f"""
            This comprehensive report presents the analysis of School-Based Distribution (SBD) data collected across Sierra Leone, 
            covering {total_districts} districts and {total_chiefdoms} chiefdoms with a total of {total_records} school records.
            
            KEY FINDINGS:
            • Total Schools Surveyed: {total_records:,}
            • Districts Covered: {total_districts}
            • Chiefdoms Covered: {total_chiefdoms}
            • Total Student Enrollment: {int(total_enrollment):,}
            • Total ITNs Distributed: {int(total_itn):,}
            • Overall Coverage Rate: {overall_coverage:.1f}%
            
            This report provides detailed analysis of enrollment patterns, ITN distribution effectiveness, 
            and geographic coverage across administrative boundaries.
            """
            doc.add_paragraph(summary_text)
            
            # Add district overview chart
            doc.add_heading('District Overview Charts', level=1)
            
            # Create district summary data for charts
            district_analysis_report = []
            for district in extracted_df['District'].dropna().unique():
                district_data = extracted_df[extracted_df['District'] == district]
                district_enrollment = 0
                district_itn = 0
                
                for class_num in range(1, 6):
                    boys_col = f"Number of boys in class {class_num}"
                    girls_col = f"Number of girls in class {class_num}"
                    itn_col = f"Number of ITN distributed to class {class_num}"
                    
                    if boys_col in district_data.columns:
                        district_enrollment += district_data[boys_col].fillna(0).sum()
                    if girls_col in district_data.columns:
                        district_enrollment += district_data[girls_col].fillna(0).sum()
                    if itn_col in district_data.columns:
                        district_itn += district_data[itn_col].fillna(0).sum()
                
                district_analysis_report.append({
                    'District': district,
                    'Enrollment': district_enrollment,
                    'ITN': district_itn,
                    'Coverage': (district_itn / district_enrollment * 100) if district_enrollment > 0 else 0
                })
            
            # Create enrollment chart
            try:
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 6))
                
                districts = [d['District'] for d in district_analysis_report]
                enrollments = [d['Enrollment'] for d in district_analysis_report]
                itns = [d['ITN'] for d in district_analysis_report]
                
                # Enrollment bar chart
                ax1.bar(districts, enrollments, color='#4682B4', edgecolor='navy')
                ax1.set_title('Total Enrollment by District', fontsize=14, fontweight='bold')
                ax1.set_ylabel('Number of Students')
                ax1.tick_params(axis='x', rotation=45)
                
                # ITN distribution chart
                ax2.bar(districts, itns, color='#32CD32', edgecolor='darkgreen')
                ax2.set_title('Total ITN Distribution by District', fontsize=14, fontweight='bold')
                ax2.set_ylabel('Number of ITNs')
                ax2.tick_params(axis='x', rotation=45)
                
                plt.tight_layout()
                
                # Save chart to temporary file
                chart_buffer = io.BytesIO()
                plt.savefig(chart_buffer, format='png', dpi=300, bbox_inches='tight')
                chart_buffer.seek(0)
                
                # Add chart to document
                doc.add_paragraph("District Performance Comparison:")
                chart_para = doc.add_paragraph()
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_run = chart_para.add_run()
                chart_run.add_picture(chart_buffer, width=Inches(6))
                
                plt.close(fig)
                
            except Exception as e:
                doc.add_paragraph(f"Chart generation note: Visual charts available in dashboard")
            
            # Add district summary table
            doc.add_heading('District Performance Summary Table', level=1)
            
            # Create district summary table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'District'
            hdr_cells[1].text = 'Schools'
            hdr_cells[2].text = 'Total Enrollment'
            hdr_cells[3].text = 'ITNs Distributed'
            hdr_cells[4].text = 'Coverage (%)'
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add district data
            for district_info in district_analysis_report:
                district = district_info['District']
                district_data = extracted_df[extracted_df['District'] == district]
                
                row_cells = table.add_row().cells
                row_cells[0].text = district
                row_cells[1].text = str(len(district_data))
                row_cells[2].text = f"{int(district_info['Enrollment']):,}"
                row_cells[3].text = f"{int(district_info['ITN']):,}"
                row_cells[4].text = f"{district_info['Coverage']:.1f}%"
            
            # Add methodology section
            doc.add_page_break()
            doc.add_heading('Methodology & Data Sources', level=1)
            
            methodology_text = f"""
            DATA COLLECTION METHODOLOGY:
            
            • Data Source: School-Based Distribution (SBD) survey data
            • Collection Period: {current_datetime.strftime('%Y')}
            • Geographic Scope: {total_districts} districts across Sierra Leone
            • Administrative Coverage: {total_chiefdoms} chiefdoms
            • School Sample Size: {total_records:,} schools
            
            VISUAL ANALYTICS INCLUDED:
            
            • Geographic Maps: District and chiefdom boundary visualizations
            • Performance Charts: Bar charts showing enrollment and ITN distribution
            • Coverage Analysis: Pie charts and performance comparisons
            • Comparative Analysis: District and chiefdom performance metrics
            
            DATA PROCESSING:
            
            • QR code extraction for geographic coordinates
            • Enrollment data aggregation across classes 1-5
            • ITN distribution tracking by class level
            • Coverage calculation: (ITNs Distributed / Total Enrollment) × 100
            • Geographic mapping using administrative boundaries (Chiefdom 2021.shp)
            """
            doc.add_paragraph(methodology_text)
            
            # Add recommendations section
            doc.add_heading('Key Recommendations', level=1)
            
            # Calculate top and bottom performing districts
            district_performance = [(d['District'], d['Coverage']) for d in district_analysis_report]
            district_performance.sort(key=lambda x: x[1], reverse=True)
            best_district = district_performance[0] if district_performance else ("N/A", 0)
            worst_district = district_performance[-1] if district_performance else ("N/A", 0)
            
            recommendations_text = f"""
            Based on the comprehensive analysis of SBD data and visual analytics, the following recommendations are proposed:
            
            IMMEDIATE ACTIONS:
            
            1. PRIORITY INTERVENTION AREAS:
               • Focus additional resources on {worst_district[0]} District (Coverage: {worst_district[1]:.1f}%)
               • Investigate supply chain issues in underperforming areas
               • Strengthen coordination with district education offices
            
            2. BEST PRACTICE REPLICATION:
               • Study successful implementation in {best_district[0]} District (Coverage: {best_district[1]:.1f}%)
               • Document and replicate effective distribution strategies
               • Share lessons learned across all districts
            
            3. MONITORING & EVALUATION:
               • Establish real-time tracking systems for ITN distribution
               • Implement quarterly review meetings with district teams
               • Develop standardized reporting formats with visual dashboards
            
            STRATEGIC RECOMMENDATIONS:
            
            • Strengthen community engagement and awareness campaigns
            • Improve supply chain management and forecasting
            • Enhance training for school-based distribution teams
            • Develop contingency plans for hard-to-reach areas
            • Establish feedback mechanisms from schools and communities
            • Implement data visualization tools for ongoing monitoring
            
            NEXT STEPS:
            
            • Conduct follow-up assessments in 6 months
            • Develop district-specific action plans based on visual analytics
            • Allocate additional resources based on performance gaps identified
            • Establish partnerships with local NGOs and community organizations
            • Create interactive dashboards for real-time monitoring
            """
            doc.add_paragraph(recommendations_text)
            
            # Add footer with generation info
            doc.add_page_break()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(f"\nComprehensive Report with Visual Analytics\nGenerated by SBD Analysis Dashboard\n{current_datetime.strftime('%B %d, %Y at %I:%M %p')}")
            footer_run.font.size = Pt(10)
            footer_run.italic = True
            
            # Save to BytesIO
            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_data = word_buffer.getvalue()
            
            st.download_button(
                label="💾 Download Complete Report with Visuals",
                data=word_data,
                file_name=f"SBD_Complete_Report_with_Visuals_{current_datetime.strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download comprehensive report with logos, charts, and maps in Word format"
            )

    # Display final summary
    st.info(f"📋 **Dataset Summary**: {len(extracted_df)} total records extracted and processed")
